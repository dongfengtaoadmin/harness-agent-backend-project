"""
RAG 领域核心：单文件聚合「分类 / 解析 / 分块 / 向量化 / 写库 / pipeline」。

教学说明：
ingest_file()               ← 唯一对外公开的函数（入口）
    │
    ├── 1. _extract_text()        解析文件文本
    ├── 2. infer_doc_category()   文件分类
    ├── 3. _split_chunks()        切分文本块
    ├── 4. _generate_embeddings() 向量化
    ├── 5. _upsert_chunks()       写入 Qdrant
    └── 返回 IngestResult
"""

import hashlib
import re
import uuid
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path

import pypdfium2 as pdfium
import docx
from docx.oxml.ns import qn
from langchain_text_splitters import RecursiveCharacterTextSplitter
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, FieldCondition, Filter, MatchValue, PointStruct, VectorParams
from sentence_transformers import SentenceTransformer

from app.core.settings import settings
from app.rag.ocr import ocr_pdf

# 业务固定常量
# ============================================================

# 免费的本地中文 Embedding 模型，无需 API Key；首次运行会下载到本机缓存。
EMBEDDING_MODEL = "BAAI/bge-small-zh-v1.5"
EMBEDDING_DIM = 512
EMBED_BATCH_SIZE = 32

# Qdrant collection 名：文件知识库现在用，对话记忆库后续接 Memory 时用。
COLLECTION_KNOWLEDGE = "knowledge_chunks"
COLLECTION_MEMORY = "conversation_memory"

# 文档分类标签。
CATEGORY_RESUME = "resume"
CATEGORY_STUDY = "study_material"
CATEGORY_GENERAL = "general"
SUPPORTED_CATEGORIES = {CATEGORY_RESUME, CATEGORY_STUDY, CATEGORY_GENERAL}

# 文件内容分类关键字：对提取后的全文做命中计数，得分高者胜。
_RESUME_CONTENT_KEYWORDS = (
    "工作经历", "项目经历", "教育背景", "个人信息", "求职意向",
    "技能", "获奖", "自我评价", "联系方式", "实习",
)
_STUDY_CONTENT_KEYWORDS = (
    "第一章", "第一节", "知识点", "总结", "练习题",
    "定义", "定理", "例题", "解题", "课程",
)

# 分块参数。
CHUNK_SIZE = 500
OVERLAP = 50
MIN_CHUNK_LEN = 20
# 递归切片分隔符：优先按标题/段落/句子切，最后才按字符兜底，比固定窗口更能保留语义结构。
_CHUNK_SEPARATORS = ["\n### ", "\n## ", "\n# ", "\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]

# 文本清洗用正则（编译一次复用）。
_INVISIBLE = re.compile(r"[\u200b\u200c\u200d\ufeff]")  # 零宽不可见字符
_MULTI_BLANK = re.compile(r"[ \t\u3000]+")              # 连续空白（含全角空格）
_MULTI_NEWLINE = re.compile(r"\n{3,}")                  # 3 个以上换行压成 2 个
_PAGE_NUMBER = re.compile(r"^\s*[-—]?\s*\d{1,4}\s*[-—]?\s*$")  # 独立成行的页码


# ============================================================
# 1) 文件分类
# ============================================================

def _infer_category_by_content(text: str) -> str:
    """按文件内容关键字计分推断分类；双方得分均为 0 则返回 general。"""
    # sum(1 for ...) 生成器计数：每命中一个关键字计 1 分。
    resume_score = sum(1 for kw in _RESUME_CONTENT_KEYWORDS if kw in text)
    study_score  = sum(1 for kw in _STUDY_CONTENT_KEYWORDS  if kw in text)
    if resume_score == 0 and study_score == 0:
        return CATEGORY_GENERAL
    # 得分相等时倾向简历（面试场景中简历更常见）。
    return CATEGORY_RESUME if resume_score >= study_score else CATEGORY_STUDY


def infer_doc_category( explicit: str | None = None, text: str = "") -> str:
    """对外公开的分类入口：显式值优先 → 内容关键字计分 → general 兜底。"""
    # SUPPORTED_CATEGORIES 白名单校验，防止脏数据写入 payload。
    if explicit and explicit in SUPPORTED_CATEGORIES:
        return explicit
    return _infer_category_by_content(text)


# ============================================================
# 2) 文件文本提取（PDF / DOCX / TXT）
# ============================================================

def _normalize(text: str) -> str:
    """统一清洗：去不可见字符 / 合并多空格 / 压缩连续换行 / 删除独立页码行，并保留段落边界。"""
    text = _INVISIBLE.sub("", text)
    lines: list[str] = []
    for ln in text.splitlines():
        stripped = ln.strip()
        # 删除独立成行的页码（PDF 抽取常见噪声），但保留空行作为段落分隔。
        if stripped and _PAGE_NUMBER.match(stripped):
            continue
        lines.append(stripped)
    text = "\n".join(lines)
    text = _MULTI_BLANK.sub(" ", text)
    text = _MULTI_NEWLINE.sub("\n\n", text)
    return text.strip()


def _parse_pdf(path: Path) -> str:
    """
    用 pypdfium2 抽取每页文本；若平均每页字符数过低，判定为扫描版并走火山 OCR 兜底。

    教学说明：电子版 PDF（文字可选中）走 pypdfium2，速度快、零成本；
    扫描版 PDF（页面是图片）才走 OCR，避免对所有 PDF 收取 OCR 成本。
    """
    pages: list[str] = []
    pdf = pdfium.PdfDocument(str(path))
    try:
        for page in pdf:
            textpage = page.get_textpage()
            t = textpage.get_text_range()
            textpage.close()
            page.close()
            if t and t.strip():
                pages.append(t.strip())
    finally:
        pdf.close()

    text = "\n".join(pages)
    page_count = max(len(pages), 1)
    avg_chars = len(text) / page_count
    print(f"[RAG][PDF] pypdfium2 抽取：页数={len(pages)}, 平均每页={avg_chars:.1f} 字")

    # 触发条件 1：完全没抽到文本（多见于纯扫描件）
    # 触发条件 2：平均每页字符数过低（可能是扫描件 + 个别页有零星可选文字）
    if not text or avg_chars < _SCANNED_PDF_AVG_CHARS:
        print("[RAG][PDF] 判定为扫描版，走火山 OCR 兜底")
        return ocr_pdf(path)

    return text


def _parse_docx(path: Path) -> str:
    """提取 Word 文本；通过段落大纲级别识别标题并注入 Markdown 标题标记，兼容中英文 Word。"""
    document = docx.Document(str(path))
    lines: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        pPr = para._p.pPr
        outline = pPr.find(qn("w:outlineLvl")) if pPr is not None else None
        outline_lvl = int(outline.get(qn("w:val"))) if outline is not None else None
        print(outline_lvl,'标题')
        if outline_lvl == 0:
            lines.append(f"# {text}")
        elif outline_lvl == 1:
            lines.append(f"## {text}")
        elif outline_lvl == 2:
            lines.append(f"### {text}")
        else:
            lines.append(text)
    return "\n".join(lines)


def _parse_txt(path: Path) -> str:
    # errors="ignore" 兼容编码不规范的文件。
    return path.read_text(encoding="utf-8", errors="ignore")


_PARSERS = {".pdf": _parse_pdf, ".docx": _parse_docx, ".txt": _parse_txt}
SUPPORTED_EXTENSIONS = frozenset(_PARSERS.keys())

# 扫描版 PDF 判定阈值：平均每页字符数 < 该值则判定为扫描版，触发火山 OCR 兜底。
_SCANNED_PDF_AVG_CHARS = 20


def _extract_text(file_path: Path) -> str:
    """按扩展名分派解析器并统一清洗。"""
    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在：{file_path}")
    # 字典分派：新增格式只需在 _PARSERS 里加一行。
    parser = _PARSERS.get(file_path.suffix.lower())
    if parser is None:
        raise ValueError(f"不支持的文件类型：{file_path.suffix}")
    return _normalize(parser(file_path))


# ============================================================
# 3) 文本分块（递归智能切片 + MD5 去重）
# ============================================================

def _filter_and_dedupe_chunks(chunks: list[str]) -> list[str]:
    """过滤过短文本块，并按 MD5 指纹去重，保留首次出现顺序。"""
    seen, unique = set(), []
    for chunk in chunks:
        chunk = chunk.strip()
        if len(chunk) < MIN_CHUNK_LEN:
            continue
        fp = hashlib.md5(chunk.encode("utf-8")).hexdigest()
        if fp not in seen:
            seen.add(fp)
            unique.append(chunk)
    return unique

# 旧方案：固定窗口 + 句子边界回退。
# 保留作教学对比：它能避免一部分硬切断，但主要仍按长度推进，无法优先感知标题/段落结构。
# def _split_chunks(text: str) -> list[str]:
#     """切分长文本：保证 start 单调推进；MD5 指纹完全相同的块去重。"""
#     if not text:
#         return []
#
#     chunks: list[str] = []
#     start, n = 0, len(text)
#     while start < n:
#         end = min(start + CHUNK_SIZE, n)
#         # 非末尾窗口尝试回退到最近的句子终止符，避免硬切断句子。
#         if end < n:
#             for i in range(end, max(start + OVERLAP, end - 100), -1):
#                 if text[i] in _SENTENCE_ENDS:
#                     end = i + 1
#                     break
#
#         chunk = text[start:end].strip()
#         if len(chunk) >= MIN_CHUNK_LEN:
#             chunks.append(chunk)
#
#         # 窗口已触底：本轮是最后一块，直接退出。
#         if end >= n:
#             break
#
#         # 推进步长 = end - overlap；若回退后没前进则强制 +1，杜绝死循环。
#         nxt = end - OVERLAP
#         start = nxt if nxt > start else start + 1
#
#     # MD5 指纹去重，保留首次出现顺序。
#     seen, unique = set(), []
#     for c in chunks:
#         fp = hashlib.md5(c.encode("utf-8")).hexdigest()
#         if fp not in seen:
#             seen.add(fp)
#             unique.append(c)
#     print(f"[RAG][分块] 切分前={len(chunks)} 块, 去重后={len(unique)} 块")
#     return unique


def _split_chunks(text: str) -> list[str]:
    """递归切分长文本：优先按标题/段落/句子边界切分；MD5 指纹完全相同的块去重。"""
    if not text:
        return []

    # 新方案：递归切片优先保留文档结构，标题/段落切不开时再逐级降级到句子、标点和字符。
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=OVERLAP,
        separators=_CHUNK_SEPARATORS,
    )
    chunks = splitter.split_text(text)
    unique = _filter_and_dedupe_chunks(chunks)
    print(f"[RAG][分块] 切分前={len(chunks)} 块, 去重后={len(unique)} 块")
    return unique


# ============================================================
# 4) 本地向量化（免费 + 模型缓存）
# ============================================================

@lru_cache(maxsize=1)
def _get_embedding_model() -> SentenceTransformer:
    """加载并缓存本地模型，避免每次摄入文档都重新初始化。"""
    print(f"[RAG][向量化] 加载本地模型：{EMBEDDING_MODEL}")
    return SentenceTransformer(EMBEDDING_MODEL)


def _generate_embeddings(texts: list[str]) -> list[list[float]]:
    """在本地批量生成归一化向量，不调用付费 API。"""
    if not texts:
        return []

    embeddings = _get_embedding_model().encode(
        texts,
        batch_size=EMBED_BATCH_SIZE,
        normalize_embeddings=True,
        show_progress_bar=False,
    )
    return embeddings.tolist()


# ============================================================
# 5) Qdrant 写入（UUID id + payload）
# ============================================================

def _get_qdrant_client() -> QdrantClient:
    return QdrantClient(host=settings.qdrant_host, port=settings.qdrant_port)


def _ensure_collection(client: QdrantClient) -> None:
    """文件知识库 collection 不存在时按维度创建，余弦相似度。"""
    if not client.collection_exists(COLLECTION_KNOWLEDGE):
        client.create_collection(
            collection_name=COLLECTION_KNOWLEDGE,
            vectors_config=VectorParams(size=EMBEDDING_DIM, distance=Distance.COSINE),
        )


def _upsert_chunks(
    chunks: list[str],
    vectors: list[list[float]],
    user_id: int,
    doc_category: str,
    file_name: str,
) -> int:
    """写向量与 payload；空块直接返回 0。"""
    if not chunks or not vectors:
        return 0
    if len(chunks) != len(vectors):
        raise ValueError("chunks 与 vectors 长度不一致")

    client = _get_qdrant_client()
    _ensure_collection(client)

    # 覆盖语义：同一个 (user_id, file_name) 重复上传时先删后写，避免新旧点并存产生检索重复。
    delete_filter = Filter(
        must=[
            FieldCondition(key="user_id", match=MatchValue(value=user_id)),
            FieldCondition(key="file_name", match=MatchValue(value=file_name)),
        ]
    )
    client.delete(collection_name=COLLECTION_KNOWLEDGE, points_selector=delete_filter)
    print(f"[RAG][写库] 已清除旧点：user_id={user_id}, file_name={file_name}")

    points = []
    for chunk_idx, (chunk_text, chunk_vector) in enumerate(zip(chunks, vectors)):
        point = PointStruct(
            id=str(uuid.uuid4()),
            vector=chunk_vector,
            payload={
                "text": chunk_text,
                "user_id": user_id,
                "doc_category": doc_category,
                "file_name": file_name,
                "chunk_index": chunk_idx,
            },
        )
        points.append(point)

    client.upsert(collection_name=COLLECTION_KNOWLEDGE, points=points)
    print(f"[RAG][写库] 写入 {len(points)} 个 point, doc_category={doc_category}")
    return len(points)


# ============================================================
# 6) 一站式摄入入口（对外唯一公开函数）
# ============================================================

@dataclass
class IngestResult:
    """摄入结果摘要：透传给上层接口。"""

    file_name: str
    doc_category: str
    chunk_count: int


def ingest_file(
    file_path: str | Path,
    user_id: int,
    doc_category: str | None = None,
    original_file_name: str | None = None,
) -> IngestResult:
    """
    单文件摄入：解析 → 清洗 → 分块去重 → 向量化 → 写 Qdrant。
    doc_category 显式合法值优先，否则按文件内容推断。
    """
    path = Path(file_path)
    file_name = original_file_name or path.name
    print(f"\n[RAG][入口] === 开始摄入 file_name={file_name}, user_id={user_id}, 显式分类={doc_category} ===")

    # 先提取文本再分类，分类器需要读取文件正文做关键字计分。
    text = _extract_text(path)
    print(f"[RAG][解析] 清洗后文本长度={len(text)} 字符")

    category = infer_doc_category(explicit=doc_category, text=text)
    print(f"[RAG][分类] 最终分类={category}")

    chunks = _split_chunks(text)
    vectors = _generate_embeddings(chunks)
    count = _upsert_chunks(
        chunks=chunks,
        vectors=vectors,
        user_id=user_id,
        doc_category=category,
        file_name=file_name,
    )
    print(f"[RAG][入口] === 摄入完成 file_name={file_name}, doc_category={category}, chunk_count={count} ===\n")
    return IngestResult(file_name=file_name, doc_category=category, chunk_count=count)
